# -*- coding: utf-8 -*-
import io
import re
import zipfile
import unicodedata
from typing import List, Tuple, Dict, Optional

import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ===============================
# --------- UTILITÁRIOS ---------
# ===============================
def no_accents_upper(s: str) -> str:
    if s is None:
        return ""
    s = unicodedata.normalize("NFKD", s).encode("ASCII", "ignore").decode("ASCII")
    return s.upper()


def group_lines(words, tol_y: float = 2.0):
    """Agrupa palavras por linha (y)."""
    linhas = []
    cur_y = None
    cur = []
    for (x0, y0, x1, y1, t) in words:
        if cur_y is None:
            cur_y = y0
            cur = [(x0, y0, x1, y1, t)]
            continue
        if abs(y0 - cur_y) <= tol_y:
            cur.append((x0, y0, x1, y1, t))
        else:
            linhas.append(sorted(cur, key=lambda z: z[0]))
            cur = [(x0, y0, x1, y1, t)]
            cur_y = y0
    if cur:
        linhas.append(sorted(cur, key=lambda z: z[0]))
    return linhas


def _append_text(base: str, extra: str) -> str:
    """Une texto cuidando de hifenização e espaços."""
    base = (base or "").rstrip()
    extra = (extra or "").lstrip()
    if base.endswith("-"):
        return base[:-1] + extra
    if base:
        return base + " " + extra
    return extra


def to_float_br(s: str) -> Optional[float]:
    if s is None:
        return None
    s = str(s).strip()
    if s == "":
        return None
    s = s.replace(".", "").replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None


# ===============================
# --- DETECÇÃO DO CABEÇALHO -----
# ===============================
END_MARKERS = [
    "DADOS ADICIONAIS",
    "INFORMACOES COMPLEMENTARES",
    "INFORMAÇÕES COMPLEMENTARES",
    "DADOS ADICIONAIS (COMPLEMENTO)",
    "INFORMACOES ADICIONAIS",
    "INFORMAÇÕES ADICIONAIS",
]

def find_header_positions(lines) -> Optional[Dict[str, float]]:
    """
    Localiza o cabeçalho da tabela e retorna x de cada coluna + y do header em "__y__".
    Alvo: COD | DESCRICAO | NCM/SH | (CST?) | CFOP | UN | QTD | V_UNITARIO | V_TOTAL
    """
    for ln in lines:
        tokens = [no_accents_upper(w[4]) for w in ln]
        if not any("NCM/SH" in t for t in tokens):
            continue
        if "CFOP" not in tokens:
            continue

        col_x = {}
        for i, (x0, y0, x1, y1, text) in enumerate(ln):
            t = no_accents_upper(text)

            # COD PROD
            if t in {"COD.", "COD", "CÓD.", "CÓD"}:
                nxt = no_accents_upper(ln[i + 1][4]) if i + 1 < len(ln) else ""
                if nxt.startswith("PROD"):
                    col_x["COD"] = x0

            # DESCRIÇÃO
            if t.startswith("DESCRICAO") or t.startswith("DESCRIÇÃO"):
                col_x["DESCRICAO"] = x0

            # NCM/SH
            if "NCM/SH" in t:
                col_x["NCM/SH"] = x0

            # CST (opcional)
            if t == "CST":
                col_x["CST"] = x0

            # CFOP
            if t == "CFOP":
                col_x["CFOP"] = x0

            # UN
            if t == "UN":
                col_x["UN"] = x0

            # QTD
            if t == "QTD":
                col_x["QTD"] = x0

            # V. UNITÁRIO / V. TOTAL (ou VALOR ...)
            if t in {"V.", "V"} and i + 1 < len(ln):
                nxt = no_accents_upper(ln[i + 1][4])
                if nxt.startswith("UNIT"):
                    col_x["V_UNITARIO"] = x0
                if nxt.startswith("TOTAL"):
                    col_x["V_TOTAL"] = x0
            if t == "VALOR" and i + 1 < len(ln):
                nxt = no_accents_upper(ln[i + 1][4])
                if nxt.startswith("UNIT"):
                    col_x["V_UNITARIO"] = x0
                if nxt.startswith("TOTAL"):
                    col_x["V_TOTAL"] = x0

        need = {"COD", "DESCRICAO", "NCM/SH", "CFOP", "UN", "QTD"}
        if len(set(col_x.keys()).intersection(need)) >= 6:
            col_x["__y__"] = ln[0][1]
            return col_x
    return None


def build_column_edges(col_x: Dict[str, float], page_width: float):
    keys_pref = ["COD", "DESCRICAO", "NCM/SH", "CST", "CFOP", "UN", "QTD", "V_UNITARIO", "V_TOTAL"]
    keys = [k for k in keys_pref if k in col_x]
    keys.sort(key=lambda k: col_x[k])
    xs = [col_x[k] for k in keys]
    edges = []
    for i, k in enumerate(keys):
        left = (xs[i - 1] + xs[i]) / 2 if i > 0 else max(0, xs[i] - 5)
        right = (xs[i] + xs[i + 1]) / 2 if i + 1 < len(xs) else page_width
        edges.append((k, left, right))
    return edges


# ===============================
# --- EXTRAÇÃO DE ITENS ----------
# ===============================
def extract_access_key_and_nf_number(file_bytes: bytes) -> Tuple[Optional[str], Optional[str], str]:
    """
    (chave_44, numero_nf_9, texto_p1)
    - Tenta pegar "CHAVE DE ACESSO" -> 44 dígitos; extrai nNF (pos 26-34 1-based => 25..34 0-based).
    - Fallback: padrão textual "Nº 000000000".
    - Retorna também o texto bruto da página 1 (para buscar PROJETO/RT).
    """
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            if len(doc) == 0:
                return None, None, ""
            txt = doc[0].get_text("text")
            upper = txt.upper()

            key = None
            pos = upper.find("CHAVE DE ACESSO")
            if pos != -1:
                tail = upper[pos:pos + 300]
                digits = re.findall(r"\d", tail)
                if len(digits) >= 44:
                    key = "".join(digits[:44])

            if not key:
                # fallback grosseiro: 44 dígitos em bloco
                m = re.search(r"\b\d[\d\s]{42,}\d\b", upper)
                if m:
                    only = re.findall(r"\d", m.group(0))
                    if len(only) >= 44:
                        key = "".join(only[:44])

            nf = None
            if key and len(key) == 44:
                nf = key25:34[1]()  # nNF (9 dígitos)

            if not nf:
                m = re.search(r"\bN[º°O\.]?\s*([0-9]{1,9})\b", upper)
                if m:
                    nf = m.group(1).zfill(9)

            return key, nf, txt
    except Exception:
        return None, None, ""


def sniff_projeto_rt(texto_primeira_pagina: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Procura 'PROJETO <algo>' e 'RT <algo>' no texto bruto.
    Retorna (projeto, rt) ou (None, None) se não achar.
    """
    up = no_accents_upper(texto_primeira_pagina)
    projeto = None
    rt = None

    mproj = re.search(r"\bPROJETO\s*([A-Z0-9\-_/\.]+)", up)
    if mproj:
        projeto = mproj.group(1).strip(" .;,")

    # RT pode aparecer como "RT: 123", "RT 123", "REQUISIÇÃO DE TRANSPORTE <RT>" etc.
    mrt = re.search(r"\bRT[:\s\-]*([A-Z0-9\-_/\.]+)", up)
    if mrt:
        rt = mrt.group(1).strip(" .;,")

    return projeto, rt


def extract_table_page(page) -> List[Dict[str, str]]:
    """
    Extrai linhas cruas (cada linha visual) mapeando palavras -> colunas via cabeçalho.
    """
    words = page.get_text("words")
    words = [(w[0], w[1], w[2], w[3], w[4]) for w in words]
    words.sort(key=lambda t: (round(t[1], 1), t[0]))
    lines = group_lines(words, tol_y=2.0)
    if not lines:
        return []

    col_x = find_header_positions(lines)
    if not col_x:
        return []

    header_y = col_x["__y__"]
    page_width = page.rect.width
    col_edges = build_column_edges(col_x, page_width)

    raw_rows = []
    past_header = False
    for ln in lines:
        y = ln[0][1]
        row_text_all_upper = " ".join(no_accents_upper(w[4]) for w in ln)

        if not past_header:
            if y <= header_y + 0.5:
                continue
            past_header = True

        if any(m in row_text_all_upper for m in END_MARKERS):
            break

        row = {k: "" for k in ["COD", "DESCRICAO", "NCM/SH", "CST", "CFOP", "UN", "QTD", "V_UNITARIO", "V_TOTAL"]}
        for (x0, y0, x1, y1, t) in ln:
            xc = (x0 + x1) / 2
            placed = False
            for (k, left, right) in col_edges:
                if left <= xc < right:
                    row[k] = _append_text(row.get(k, ""), t)
                    placed = True
                    break
            if not placed:
                row["DESCRICAO"] = _append_text(row.get("DESCRICAO", ""), t)

        for k in row:
            row[k] = " ".join(str(row[k]).replace("\n", " ").split())
        raw_rows.append(row)

    return raw_rows


def is_new_item(row: Dict[str, str]) -> bool:
    """Novo item quando a linha traz COD e NCM/SH visivelmente preenchidos."""
    return bool(row.get("COD", "").strip()) and bool(row.get("NCM/SH", "").strip())


def parse_ncm_cst_cfop(ncm_text: str, cst_text: str, cfop_text: str) -> Tuple[str, str, str]:
    """
    (NCM8, CST3, CFOP4) a partir de textos crus — ordem: 8 -> 3 -> 4 (com fallbacks).
    """
    def tokens_digits(s: str) -> List[str]:
        return [t for t in re.split(r"\D+", s or "") if t]

    toks = tokens_digits((ncm_text or "") + " " + (cst_text or "") + " " + (cfop_text or ""))

    ncm = cst = cfop = ""
    i = 0
    while i < len(toks):
        if len(toks[i]) == 8:
            ncm = toks[i]; i += 1; break
        i += 1
    while i < len(toks):
        if len(toks[i]) == 3:
            cst = toks[i]; i += 1; break
        i += 1
    while i < len(toks):
        if len(toks[i]) == 4:
            cfop = toks[i]; i += 1; break
        i += 1

    if not ncm:
        m = re.search(r"\b(\d{8})\b", ncm_text or "")
        if m: ncm = m.group(1)
    if not cst:
        m = re.search(r"\b(\d{3})\b", cst_text or "")
        if m: cst = m.group(1)
    if not cfop:
        m = re.search(r"\b(\d{4})\b", cfop_text or "")
        if m: cfop = m.group(1)

    return ncm, cst, cfop


def consolidate_rows_into_items(raw_rows: List[Dict[str, str]]) -> List[Dict[str, str]]:
    """
    Consolida N linhas em 1 item (descrição unificada).
    """
    final_rows = []
    current = None

    for r in raw_rows:
        if is_new_item(r):
            if current:
                ncm, cst, cfop = parse_ncm_cst_cfop(current.get("NCM/SH", ""), current.get("CST", ""), current.get("CFOP", ""))
                current["NCM/SH"], current["CST"], current["CFOP"] = ncm, cst, cfop
                final_rows.append(current)

            current = {k: r.get(k, "").strip() for k in ["COD", "DESCRICAO", "NCM/SH", "CST", "CFOP", "UN", "QTD", "V_UNITARIO", "V_TOTAL"]}
            continue

        # Continuação
        if current:
            if r.get("DESCRICAO", "").strip():
                current["DESCRICAO"] = _append_text(current["DESCRICAO"], r["DESCRICAO"])
            for col in ["CST", "CFOP", "UN", "QTD", "V_UNITARIO", "V_TOTAL", "NCM/SH"]:
                if not current.get(col, "").strip() and r.get(col, "").strip():
                    current[col] = r[col].strip()

    if current:
        ncm, cst, cfop = parse_ncm_cst_cfop(current.get("NCM/SH", ""), current.get("CST", ""), current.get("CFOP", ""))
        current["NCM/SH"], current["CST"], current["CFOP"] = ncm, cst, cfop
        final_rows.append(current)

    # limpeza final
    for r in final_rows:
        for k in r:
            r[k] = " ".join(str(r[k]).split())

    return final_rows


def extract_items_from_pdf(file_bytes: bytes) -> pd.DataFrame:
    """Extrai itens (linha única por item) desta DANFE."""
    out_rows = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for p in doc:
            raw_rows = extract_table_page(p)
            out_rows.extend(raw_rows)
    items = consolidate_rows_into_items(out_rows)
    df = pd.DataFrame(items, columns=["COD", "DESCRICAO", "NCM/SH", "CST", "CFOP", "UN", "QTD", "V_UNITARIO", "V_TOTAL"])
    return df


# ===============================
# --- CAMPOS MÁSCARA 1–23 -------
# ===============================
MASK_FIELDS = [
    "1-", "2-DOCUMENTO DE CHEGADA", "3-MATERIAL", "4-NOTA PETROBRAS", "5-NOTA DE SAÍDA BASE",
    "6-NOTA DE SAÍDA PETROBRAS", "7-N° DE CSP", "8-PROJETO", "9-RT", "10-MATERIAL INVENTARIADO",
    "11-FERRAMENTA", "12-TAG BCDS DA FERRAMENTA", "13-NM", "14-CENTRO", "15-DESENHO",
    "16-IMOBILIZADO", "17-QUANTIDADE", "18-N° DE CAIXA", "19-DIAGRAMA DE REDE / ELEMENTO PEP",
    "20-PTM", "21-DSM", "22-NOTA DE TRANSF. MAR", "23-PROTOCOLO"
]

RE_IT = re.compile(r"\bIT\s*\d{2,}\b", re.IGNORECASE)
RE_NM = re.compile(r"\bNM\.?\s*(\d{5,})\b", re.IGNORECASE)

def extrair_item_unifilar(descricao: str) -> str:
    m = RE_IT.search(descricao or "")
    return m.group(0).upper().replace(" ", "") if m else ""

def extrair_nm(descricao: str) -> str:
    m = RE_NM.search(descricao or "")
    return m.group(1) if m else ""


def build_mask_defaults(nf_num: str,
                        chave: str,
                        projeto_hint: Optional[str],
                        rt_hint: Optional[str],
                        item_row: pd.Series) -> Dict[str, str]:
    """
    Gera um dicionário com os 23 campos já pré-preenchidos quando possível.
    """
    desc = str(item_row.get("DESCRICAO", "") or "")
    nm = extrair_nm(desc)
    desenho = str(item_row.get("COD", "") or "")
    qtd = str(item_row.get("QTD", "") or "")
    vunit = str(item_row.get("V_UNITARIO", "") or "")

    # Você pode ajustar os defaults abaixo conforme seu processo
    defaults = {
        "1-": "",
        "2-DOCUMENTO DE CHEGADA": f"DANFE NF {nf_num}" if nf_num else "DANFE",
        "3-MATERIAL": desc,
        "4-NOTA PETROBRAS": nf_num or "",
        "5-NOTA DE SAÍDA BASE": "N/A",
        "6-NOTA DE SAÍDA PETROBRAS": "N/A",
        "7-N° DE CSP": "",
        "8-PROJETO": projeto_hint or "",
        "9-RT": rt_hint or "N/A",
        "10-MATERIAL INVENTARIADO": "",
        "11-FERRAMENTA": "NÃO",
        "12-TAG BCDS DA FERRAMENTA": "N/A",
        "13-NM": nm,
        "14-CENTRO": "",
        "15-DESENHO": desenho if desenho else "N/A",
        "16-IMOBILIZADO": "N/A",
        "17-QUANTIDADE": qtd,
        "18-N° DE CAIXA": "N/A",
        "19-DIAGRAMA DE REDE / ELEMENTO PEP": "N/A",
        "20-PTM": "N/A",
        "21-DSM": "N/A",
        "22-NOTA DE TRANSF. MAR": "N/A",
        "23-PROTOCOLO": "N/A",
    }
    return defaults


def render_mask_docx(mask_dict: Dict[str, str]) -> bytes:
    """
    Gera um DOCX com os 23 campos numerados.
    """
    doc = Document()
    # Título
    title = doc.add_paragraph("MÁSCARA DE ENVIO — ITENS NF")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    doc.add_paragraph("")  # espaço

    for field in MASK_FIELDS:
        value = mask_dict.get(field, "")
        p = doc.add_paragraph()
        run1 = p.add_run(f"{field} ")
        run1.bold = True
        run1.font.size = Pt(11)

        run2 = p.add_run(value)
        run2.font.size = Pt(11)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def render_mask_xlsx(mask_dict: Dict[str, str]) -> bytes:
    """
    Gera um XLSX com uma linha e as 23 colunas (campos).
    """
    df = pd.DataFrame([{k: mask_dict.get(k, "") for k in MASK_FIELDS}])
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as w:
        df.to_excel(w, index=False, sheet_name="Mascara")
    out.seek(0)
    return out.getvalue()


# ===============================
# --------- INTERFACE -----------
# ===============================
st.set_page_config(page_title="Coletor de Campos da NF + Geração de Máscara", layout="wide")
st.title("🧾 Coletor de NFs → Preenchimento de Máscara (1–23)")

st.markdown(
    """
**Fluxo sugerido**  
1. Faça **upload** de **uma ou mais DANFEs (PDF)**;  
2. O app extrai **itens consolidados** (uma linha por item) e **dados do cabeçalho** (Nº NF, Chave, PROJETO/RT se houver);  
3. Selecione o **item** e revise/edite os **23 campos**;  
4. **Baixe a máscara** preenchida em **DOCX** e/ou **XLSX** (ou gere em **lote** para vários itens).
"""
)

files = st.file_uploader("Selecione uma ou mais DANFEs (PDF)", type=["pdf"], accept_multiple_files=True)

if files:
    # Extração de todas as NFs
    all_rows = []
    per_nf_meta = {}  # nf_number -> (chave, projeto, rt)

    for f in files:
        fbytes = f.read()
        chave, nf_num, txt_p1 = extract_access_key_and_nf_number(fbytes)
        projeto_hint, rt_hint = sniff_projeto_rt(txt_p1)
        nf_label = nf_num if nf_num else "NF_DESCONHECIDA"

        df_items = extract_items_from_pdf(fbytes)
        if df_items.empty:
            st.warning(f"⚠️ Nenhum item encontrado: {f.name}")
            continue

        df_items.insert(0, "NF", nf_label)
        df_items.insert(1, "Arquivo", f.name)
        df_items.insert(2, "Chave_Acesso", chave or "")

        all_rows.append(df_items)
        per_nf_meta[nf_label] = (chave or "", nf_label, projeto_hint, rt_hint)

    if not all_rows:
        st.stop()

    df_all = pd.concat(all_rows, ignore_index=True)

    st.success(f"NF(s) processadas: {df_all['NF'].nunique()} • Itens extraídos: {len(df_all)}")
    st.dataframe(df_all, use_container_width=True, height=420)

    # ===============================
    # Seleção de item único
    # ===============================
    st.markdown("### Selecionar **um item** para preencher a máscara")
    colNF, colIdx = st.columns([1, 2])

    with colNF:
        sel_nf = st.selectbox("NF", sorted(df_all["NF"].unique()))
    with colIdx:
        # opções do item daquela NF (mostra uma string amigável)
        df_nf = df_all[df_all["NF"] == sel_nf].reset_index(drop=True)
        opts = [
            f"#{i+1} • COD={row['COD']} • QTD={row['QTD']} • UN={row['UN']} • DESC={row['DESCRICAO'][:60]}..."
            for i, row in df_nf.iterrows()
        ]
        sel_idx = st.selectbox("Item da NF", options=list(range(len(df_nf))), format_func=lambda i: opts[i])

    # Pré-preenche máscara
    chave, nf_num, projeto_hint, rt_hint = per_nf_meta.get(sel_nf, ("", sel_nf, None, None))
    default_mask = build_mask_defaults(nf_num, chave, projeto_hint, rt_hint, df_nf.loc[sel_idx])

    st.markdown("#### Revisar/editar campos da **máscara**")
    # Formulário em duas colunas
    form_cols = st.columns(2)
    mask_vals = {}
    for i, field in enumerate(MASK_FIELDS):
        col = form_cols[i % 2]
        with col:
            mask_vals[field] = st.text_input(field, value=default_mask.get(field, ""))

    # Botões de saída
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("📄 Gerar DOCX (este item)"):
            doc_bytes = render_mask_docx(mask_vals)
            st.download_button("⬇️ Baixar DOCX", data=doc_bytes, file_name=f"mascara_{sel_nf}_{df_nf.loc[sel_idx,'COD']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with c2:
        if st.button("📊 Gerar XLSX (este item)"):
            xls_bytes = render_mask_xlsx(mask_vals)
            st.download_button("⬇️ Baixar XLSX", data=xls_bytes, file_name=f"mascara_{sel_nf}_{df_nf.loc[sel_idx,'COD']}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    st.markdown("---")

    # ===============================
    # Geração em LOTE (vários itens)
    # ===============================
    st.markdown("### Geração **em lote** (vários itens → ZIP)")
    # Filtros de lote
    lote_nf = st.multiselect("Filtrar por NF (opcional)", options=sorted(df_all["NF"].unique()), default=sorted(df_all["NF"].unique()))
    df_lote = df_all[df_all["NF"].isin(lote_nf)].reset_index(drop=True)
    st.caption(f"Itens selecionáveis: {len(df_lote)}")
    sel_rows = st.multiselect(
        "Escolha os itens para gerar em lote (por índice da tabela abaixo)",
        options=list(df_lote.index),
        format_func=lambda i: f"[{i}] NF={df_lote.loc[i,'NF']} • COD={df_lote.loc[i,'COD']} • QTD={df_lote.loc[i,'QTD']} • {df_lote.loc[i,'DESCRICAO'][:40]}..."
    )
    st.dataframe(df_lote, use_container_width=True, height=320)

    czip1, czip2 = st.columns(2)
    with czip1:
        if st.button("🧷 Gerar ZIP de **DOCX**"):
            if not sel_rows:
                st.warning("Selecione ao menos um item.")
            else:
                zip_buf = io.BytesIO()
                with zipfile.ZipFile(zip_buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
                    for i in sel_rows:
                        row = df_lote.loc[i]
                        chave, nf_num, projeto_hint, rt_hint = per_nf_meta.get(row["NF"], ("", row["NF"], None, None))
                        defaults = build_mask_defaults(nf_num, chave, projeto_hint, rt_hint, row)
                        doc_bytes = render_mask_docx(defaults)
                        fname = f"mascara_{row['NF']}_{row['COD']}.docx"
                        zf.writestr(fname, doc_bytes)
                zip_buf.seek(0)
                st.download_button("⬇️ Baixar ZIP (DOCX)", data=zip_buf.getvalue(), file_name="mascaras_docx.zip", mime="application/zip")
    with czip2:
        if st.button("🧷 Gerar ZIP de **XLSX**"):
            if not sel_rows:
                st.warning("Selecione ao menos um item.")
            else:
                zip_buf = io.BytesIO()
                with zipfile.ZipFile(zip_buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
                    for i in sel_rows:
                        row = df_lote.loc[i]
                        chave, nf_num, projeto_hint, rt_hint = per_nf_meta.get(row["NF"], ("", row["NF"], None, None))
                        defaults = build_mask_defaults(nf_num, chave, projeto_hint, rt_hint, row)
                        xls_bytes = render_mask_xlsx(defaults)
                        fname = f"mascara_{row['NF']}_{row['COD']}.xlsx"
                        zf.writestr(fname, xls_bytes)
                zip_buf.seek(0)
                st.download_button("⬇️ Baixar ZIP (XLSX)", data=zip_buf.getvalue(), file_name="mascaras_xlsx.zip", mime="application/zip")
else:
    st.info("Envie ao menos uma DANFE (PDF) para iniciar.")
